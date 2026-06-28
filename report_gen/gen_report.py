"""
Sinh báo cáo CNPM MetBusTrip → Mau Cuon Bao Cao Do An CNPM.docx
Dùng: python-docx
"""
import os, sys

# === Đường dẫn ===
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_FILE = os.path.join(
    os.path.dirname(BASE_DIR),
    "Mau Cuon Bao Cao Do An CNPM.docx"
)

# === Cài thư viện nếu chưa có ===
def ensure_pkg(name):
    try:
        __import__(name)
    except ImportError:
        import subprocess
        subprocess.run(
            [sys.executable, "-m", "pip", "install", name, "-q"],
            check=True
        )

ensure_pkg("python-docx")
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === Import data ===
sys.path.insert(0, BASE_DIR)
from report_data import COVER, LOI_NOI_DAU, CHUONG_1, CHUONG_2
from report_data_ch3 import CHUONG_3, ACTORS, USE_CASES, NON_FUNC_REQUIREMENTS
from report_data_ch456 import CHUONG_4, CHUONG_5, CHUONG_6, TAI_LIEU_THAM_KHAO

# === Helpers ===
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "2E75B6")
        tblBorders.append(b)
    tblPr.append(tblBorders)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def sub_heading(doc, text):
    """Heading cấp 4 (bold, italic, màu riêng)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(4)
    return p


def body_para(doc, text, indent=False):
    # Nếu text bắt đầu bằng số (VD: "5.4.1." hoặc "1." hay "a."),
    # coi là sub-heading cấp 3
    stripped = text.strip()
    if stripped and stripped[0].isdigit():
        return sub_heading(doc, stripped)
    style = "List Bullet" if text.startswith("·") or indent else "Normal"
    p = doc.add_paragraph(style=style)
    if text.startswith("·"):
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
    return p

def add_table(doc, header, rows, title=None):
    if title:
        p = doc.add_paragraph(title)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    set_cell_borders(table)

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(header):
        cell = hdr.cells[i]
        set_cell_bg(cell, "2E75B6")
        p = cell.paragraphs[0]
        p.add_run(h).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = "Times New Roman"

    # Data rows
    for row in rows:
        tr = table.add_row()
        for i, val in enumerate(row):
            cell = tr.cells[i]
            p = cell.paragraphs[0]
            p.add_run(str(val))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

    doc.add_paragraph()
    return table

def add_image(doc, rel_path, width=Cm(15), caption=None):
    full = os.path.join(BASE_DIR, rel_path)
    if not os.path.exists(full):
        p = doc.add_paragraph(f"[Hình: {rel_path} - chưa có file]")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        return
    try:
        doc.add_picture(full, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.italic = True
            cap.runs[0].font.size = Pt(10)
    except Exception as e:
        p = doc.add_paragraph(f"[Lỗi chèn ảnh {rel_path}: {e}]")

def uc_table(doc):
    """Bảng Use Case phân loại theo actor."""
    cols = 4
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    set_cell_borders(table)

    hdr = table.rows[0]
    for i, h in enumerate(["Mã UC", "Tên Use Case", "Actor", "Mô tả"]):
        cell = hdr.cells[i]
        set_cell_bg(cell, "2E75B6")
        p = cell.paragraphs[0]
        p.add_run(h).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = "Times New Roman"

    for uc in USE_CASES:
        tr = table.add_row()
        for i, val in enumerate(uc[:4]):
            cell = tr.cells[i]
            p = cell.paragraphs[0]
            p.add_run(str(val))
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

    doc.add_paragraph()


# ============================================================
# BUILD DOCUMENT
# ============================================================
def build():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ============================
    # TRANG BÌA
    # ============================
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph(COVER["truong"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = "Times New Roman"

    p = doc.add_paragraph(COVER["khoa"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.name = "Times New Roman"

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph(COVER["de_tai"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.name = "Times New Roman"

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph(f"Môn: {COVER['mon']}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    p = doc.add_paragraph(f"GVHD: {COVER['gvhd']}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    doc.add_paragraph()
    for sv in COVER["sv_thuc_hien"]:
        p = doc.add_paragraph(f"{sv[2]}: {sv[0]}  –  MSSV: {sv[1]}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(COVER["nam"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.name = "Times New Roman"

    doc.add_page_break()

    # ============================
    # MỤC LỤC (ghi tay)
    # ============================
    heading(doc, "MỤC LỤC", 1)
    toc_items = [
        ("", "Trang"),
        ("Lời nói đầu", "3"),
        ("Chương 1. Giới thiệu đề tài", "4"),
        ("Chương 2. Cơ sở lý thuyết và công nghệ", "6"),
        ("Chương 3. Phân tích yêu cầu hệ thống", "9"),
        ("Chương 4. Thiết kế hệ thống", "14"),
        ("Chương 5. Cài đặt và kết quả", "20"),
        ("Chương 6. Kết luận và hướng phát triển", "24"),
        ("Tài liệu tham khảo", "26"),
    ]
    add_table(doc,
        ["Nội dung", "Trang"],
        [[t, p] for t, p in toc_items if t]
    )

    doc.add_page_break()

    # ============================
    # LỜI NÓI ĐẦU
    # ============================
    heading(doc, "LỜI NÓI ĐẦU", 1)
    for para in LOI_NOI_DAU:
        body_para(doc, para)
    doc.add_page_break()

    # ============================
    # CHƯƠNG 1
    # ============================
    heading(doc, CHUONG_1["tieu_de"], 1)
    for muc in CHUONG_1["cac_muc"]:
        heading(doc, muc["ten"], 2)
        if "noi_dung" in muc:
            for nd in muc["noi_dung"]:
                body_para(doc, nd)
        if "bang" in muc:
            b = muc["bang"]
            add_table(doc, b["header"], b["rows"], b.get("tieu_de"))

    doc.add_page_break()

    # ============================
    # CHƯƠNG 2
    # ============================
    heading(doc, CHUONG_2["tieu_de"], 1)
    for muc in CHUONG_2["cac_muc"]:
        heading(doc, muc["ten"], 2)
        for nd in muc["noi_dung"]:
            body_para(doc, nd)

    doc.add_page_break()

    # ============================
    # CHƯƠNG 3
    # ============================
    heading(doc, CHUONG_3["tieu_de"], 1)
    for muc in CHUONG_3["cac_muc"]:
        heading(doc, muc["ten"], 2)
        if "noi_dung" in muc:
            for nd in muc["noi_dung"]:
                body_para(doc, nd)
        if "bang" in muc:
            b = muc["bang"]
            add_table(doc, b["header"], b["rows"], b.get("tieu_de"))
        if "hinh" in muc:
            add_image(doc, f"diagrams/{muc['hinh']}",
                      caption=f"[Sơ đồ: {muc['hinh']}]")

    doc.add_page_break()

    # ============================
    # CHƯƠNG 4
    # ============================
    heading(doc, CHUONG_4["tieu_de"], 1)
    for muc in CHUONG_4["cac_muc"]:
        heading(doc, muc["ten"], 2)
        if "noi_dung" in muc:
            for nd in muc["noi_dung"]:
                body_para(doc, nd)
        if "bang" in muc:
            b = muc["bang"]
            add_table(doc, b["header"], b["rows"], b.get("tieu_de"))
        if "hinh" in muc:
            add_image(doc, f"diagrams/{muc['hinh']}",
                      caption=f"[Sơ đồ: {muc['hinh']}]")

    doc.add_page_break()

    # ============================
    # CHƯƠNG 5
    # ============================
    heading(doc, CHUONG_5["tieu_de"], 1)
    for muc in CHUONG_5["cac_muc"]:
        heading(doc, muc["ten"], 2)
        if "noi_dung" in muc:
            for nd in muc["noi_dung"]:
                body_para(doc, nd)
        if "bang" in muc:
            b = muc["bang"]
            add_table(doc, b["header"], b["rows"], b.get("tieu_de"))

    doc.add_page_break()

    # ============================
    # CHƯƠNG 6
    # ============================
    heading(doc, CHUONG_6["tieu_de"], 1)
    for muc in CHUONG_6["cac_muc"]:
        heading(doc, muc["ten"], 2)
        for nd in muc["noi_dung"]:
            body_para(doc, nd)

    doc.add_page_break()

    # ============================
    # TÀI LIỆU THAM KHẢO
    # ============================
    heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    for i, ref in enumerate(TAI_LIEU_THAM_KHAO, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(ref)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"

    # ============================
    # SAVE
    # ============================
    os.makedirs(os.path.dirname(OUT_FILE), exist_ok=True)
    doc.save(OUT_FILE)
    print(f"Da luu: {OUT_FILE}")

if __name__ == "__main__":
    build()
